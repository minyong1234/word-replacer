from docx import Document
from docx.shared import RGBColor
import re

def replace_in_docx(input_path, output_path, replacements, case_sensitive=True):
    doc = Document(input_path)

    def replace_run(run):
        for old, new in replacements.items():
            if case_sensitive:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                if re.search(re.escape(old), run.text, flags=re.IGNORECASE):
                    run.text = re.sub(re.escape(old), new, run.text, flags=re.IGNORECASE)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    def replace_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                replace_run(run)

    def replace_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_paragraphs(cell.paragraphs)
                    replace_tables(cell.tables)

    # 일반 단락
    replace_paragraphs(doc.paragraphs)

    # 표
    replace_tables(doc.tables)

    # 머릿글 / 바닥글
    for section in doc.sections:
        # 머릿글
        replace_paragraphs(section.header.paragraphs)
        replace_tables(section.header.tables)

        # 바닥글
        replace_paragraphs(section.footer.paragraphs)
        replace_tables(section.footer.tables)

        # 첫 페이지 머릿글/바닥글 (설정된 경우)
        if section.different_first_page_header_footer:
            replace_paragraphs(section.first_page_header.paragraphs)
            replace_tables(section.first_page_header.tables)
            replace_paragraphs(section.first_page_footer.paragraphs)
            replace_tables(section.first_page_footer.tables)

    doc.save(output_path)
