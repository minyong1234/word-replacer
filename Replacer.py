from docx import Document
import re

def replace_in_docx(input_path, output_path, replacements, case_sensitive=True):
    doc = Document(input_path)

    def replace_text(text):
        for old, new in replacements.items():
            if case_sensitive:
                text = text.replace(old, new)
            else:
                text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
        return text

    # 일반 단락 처리 (서식 유지)
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = replace_text(run.text)

    # 표 안의 텍스트 처리
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = replace_text(run.text)

    doc.save(output_path)
